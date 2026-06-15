# -*- coding: utf-8 -*-
"""
Сборка главы 3 в Word (.docx).

Оформление текста: Times New Roman 14, поля 3/1.5/2/2 см, межстрочный 1.5,
первая строка 1.25 см, по ширине, номер страницы внизу по центру.

Таблицы (по методичке, разд. 9): сквозная нумерация «Таблица N – Название»
(без точки в конце), подпись слева без абзацного отступа; ссылка в тексте до
таблицы; рамка слева, справа и снизу; шрифт внутри таблицы 12 пт.

Запуск: py -3 docs\\build_glava3_docx.py
"""

from __future__ import annotations
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent))
from glava3_content import CHAPTER_INTRO, SEC_31, SEC_32, SEC_33, SEC_34

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "GLAVA3_struktura_i_realizatsiya.docx"

TABLE_FONT_PT = 12
_table_num = 0


def _apply_margins(doc: Document) -> None:
    sec = doc.sections[0]
    sec.left_margin = Cm(3.0)
    sec.right_margin = Cm(1.5)
    sec.top_margin = Cm(2.0)
    sec.bottom_margin = Cm(2.0)


def _apply_normal_style(doc: Document) -> None:
    st = doc.styles["Normal"]
    st.font.name = "Times New Roman"
    st.font.size = Pt(14)
    pf = st.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.first_line_indent = Cm(1.25)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.5
    pf.space_after = Pt(0)


def _add_page_number_footer(doc: Document) -> None:
    sec = doc.sections[0]
    foot = sec.footer
    p = foot.paragraphs[0] if foot.paragraphs else foot.add_paragraph()
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld1 = OxmlElement("w:fldChar")
    fld1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld1, instr, fld2])
    for r in p.runs:
        r.font.name = "Times New Roman"
        r.font.size = Pt(14)


def _p(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.5
    for r in p.runs:
        r.font.name = "Times New Roman"
        r.font.size = Pt(14)


def _ps(doc: Document, paragraphs: list[str]) -> None:
    for t in paragraphs:
        _p(doc, t)


def _h(doc: Document, text: str, level: int) -> None:
    doc.add_heading(text, level=level)
    p = doc.paragraphs[-1]
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(18 if level <= 1 else 12)
    p.paragraph_format.space_after = Pt(12 if level <= 1 else 6)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 0 else WD_ALIGN_PARAGRAPH.JUSTIFY
    for r in p.runs:
        r.font.name = "Times New Roman"
        r.bold = True
        r.font.size = Pt(14)


def _border_element(name: str, val: str) -> OxmlElement:
    el = OxmlElement(f"w:{name}")
    el.set(qn("w:val"), val)
    el.set(qn("w:sz"), "4")
    el.set(qn("w:space"), "0")
    el.set(qn("w:color"), "auto")
    return el


def _apply_table_borders(table) -> None:
    """Рамка таблицы: слева, справа, снизу; внутренние горизонтальные линии."""
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)
    borders = OxmlElement("w:tblBorders")
    spec = {
        "top": "nil",
        "left": "single",
        "bottom": "single",
        "right": "single",
        "insideH": "single",
        "insideV": "nil",
    }
    for edge, val in spec.items():
        borders.append(_border_element(edge, val))
    old = tbl_pr.find(qn("w:tblBorders"))
    if old is not None:
        tbl_pr.remove(old)
    tbl_pr.append(borders)


def _style_table_cell_text(table, *, bold_header: bool = True) -> None:
    for ri, row in enumerate(table.rows):
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.first_line_indent = Cm(0)
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for r in p.runs:
                    r.font.name = "Times New Roman"
                    r.font.size = Pt(TABLE_FONT_PT)
                    if bold_header and ri == 0:
                        r.bold = True


def _table_caption(doc: Document, number: int, title: str) -> None:
    """Подпись: «Таблица N – Название», слева, без отступа, без точки в конце."""
    p = doc.add_paragraph(f"Таблица {number} – {title}")
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for r in p.runs:
        r.font.name = "Times New Roman"
        r.font.size = Pt(14)
        r.bold = False
        r.italic = False


def _add_numbered_table(doc: Document, title: str, cols: int, header: list[str], rows: list[tuple]) -> int:
    global _table_num
    _table_num += 1
    num = _table_num
    _table_caption(doc, num, title)
    t = doc.add_table(rows=1, cols=cols)
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(header):
        hdr[i].text = h
    for row_data in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row_data):
            cells[i].text = str(val)
    _apply_table_borders(t)
    _style_table_cell_text(t)
    doc.add_paragraph()
    return num


def add_dir_table(doc: Document) -> None:
    _add_numbered_table(
        doc,
        "Назначение основных файлов проекта ppo_crypto_agent",
        2,
        ["Путь", "Назначение"],
        [
            ("envs/crypto_env.py", "Среда Gymnasium: шаг, награда, funding, ограничение |Δw|"),
            ("utils/data.py", "Чтение CSV, build_features, маска valid"),
            ("utils/inference.py", "Сборка наблюдения и apply_max_delta вне среды"),
            ("scripts/train_ppo.py", "Обучение PPO, EvalCallback, сохранение чекпоинтов"),
            ("scripts/eval_ppo.py", "Оценка политики: SB3 + buy-and-hold + детальный прогон"),
            ("scripts/live_signal.py", "Сигнал w по последним свечам, без ордеров"),
            ("scripts/make_sample_csv.py", "Синтетический ряд sample_ohlcv.csv"),
            ("scripts/merge_binance_funding.py", "Подтягивание funding Binance в CSV"),
            ("data/", "Входные наборы OHLCV"),
            ("runs/<ран>/", "best/best_model.zip, ppo_final.zip, logs/, tb/"),
            ("requirements.txt", "Фиксация версий зависимостей"),
            ("INSTALL_WINDOWS.txt", "Порядок установки и примеры команд (Windows)"),
        ],
    )


def add_hyperparam_table(doc: Document) -> None:
    _add_numbered_table(
        doc,
        "Основные гиперпараметры PPO в train_ppo.py",
        3,
        ["Параметр", "Значение", "Смысл"],
        [
            ("Политика", "MlpPolicy", "Две сети: политика π(a|s) и критик V(s)"),
            ("learning_rate", "1e-4", "Скорость обновления весов Adam"),
            ("n_steps", "2048", "Длина буфера траекторий на одно обновление"),
            ("batch_size", "256", "Размер минибатча SGD"),
            ("n_epochs", "10", "Число эпох на буфере за итерацию"),
            ("gamma", "0.99", "Дисконт будущей награды"),
            ("gae_lambda", "0.95", "Компромисс смещение/дисперсия в GAE"),
            ("clip_range", "0.2", "Отсечение отношения π_new/π_old"),
            ("ent_coef", "0.005", "Вес энтропийного бонуса (исследование)"),
            ("vf_coef", "0.5", "Вес ошибки функции ценности"),
            ("max_grad_norm", "0.5", "Клиппинг нормы градиента"),
            ("ortho_init", "True", "Ортогональная инициализация слоёв"),
        ],
    )


def build() -> None:
    global _table_num
    _table_num = 0

    doc = Document()
    _apply_margins(doc)
    _apply_normal_style(doc)
    _add_page_number_footer(doc)

    _h(doc, "Глава 3. Структура программной реализации, методика и обоснование решений", 0)
    _ps(doc, CHAPTER_INTRO)

    _h(doc, "3.1. Цель и границы задачи, архитектура решения, структура каталогов", 1)
    _ps(doc, SEC_31)
    add_dir_table(doc)

    _h(doc, "3.2. Данные, признаки, причинность, разбиение train/validation", 1)
    _ps(doc, SEC_32)

    _h(doc, "3.3. Марковская среда (MDP): действие, наблюдение, переход, награда, эпизод", 1)
    _ps(doc, SEC_33)

    _h(doc, "3.4. Обучение PPO, оценка, инференс, ограничения и направления развития", 1)
    _ps(doc, SEC_34[:4])
    _p(
        doc,
        "Основные гиперпараметры обучения PPO, зафиксированные в скрипте train_ppo.py, "
        "приведены в таблице 2.",
    )
    add_hyperparam_table(doc)
    _ps(doc, SEC_34[4:])

    all_text = "\n".join(CHAPTER_INTRO + SEC_31 + SEC_32 + SEC_33 + SEC_34)
    print(f"Tekst: ~{len(all_text)} znakovov ~ {len(all_text)/2800:.1f} str. (bez tablits)")

    try:
        doc.save(OUT)
        print(f"Zapisano: {OUT}")
    except PermissionError:
        alt = OUT.with_name(OUT.stem + "_копия.docx")
        doc.save(alt)
        print(f"Kopiya: {alt}")


if __name__ == "__main__":
    build()
